"""
Ingestion service – pluggable multi-format document loader.

Improvements implemented:
  ✅ Multi-format loaders: DOCX, HTML, Markdown, and plain text,
     each behind a unified load(file) → list[Document] interface.
  ✅ Auto-detection by file extension routes to the correct adapter.
"""

from __future__ import annotations

import re
from pathlib import Path
from langchain_core.documents import Document



def _load_docx(file) -> list[Document]:
    """Extract paragraphs from a .docx file using python-docx."""
    try:
        import docx  # python-docx
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    doc = docx.Document(file)
    name = getattr(file, "name", "document.docx")
    pages: list[Document] = []
    buffer: list[str] = []
    para_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        buffer.append(text)
        para_num += 1
        # Group every 20 paragraphs into one "page" chunk
        if para_num % 20 == 0:
            pages.append(Document(
                page_content="\n".join(buffer),
                metadata={"source": name, "page": para_num // 20, "format": "docx"},
            ))
            buffer = []

    if buffer:
        pages.append(Document(
            page_content="\n".join(buffer),
            metadata={"source": name, "page": (para_num // 20) + 1, "format": "docx"},
        ))
    return pages


def _load_html(file) -> list[Document]:
    """Strip HTML tags and return cleaned text as a single Document."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("Install beautifulsoup4: pip install beautifulsoup4")

    raw = file.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")

    soup = BeautifulSoup(raw, "html.parser")
    # Remove script/style noise
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    name = getattr(file, "name", "page.html")

    return [Document(page_content=text, metadata={"source": name, "page": 1, "format": "html"})]


def _load_markdown(file) -> list[Document]:
    """Split Markdown on H2/H3 headings into separate Documents."""
    raw = file.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")

    name = getattr(file, "name", "doc.md")
    # Split on ## or ### headings
    sections = re.split(r"(?m)^#{2,3}\s+", raw)
    docs: list[Document] = []
    for i, section in enumerate(sections, start=1):
        section = section.strip()
        if section:
            docs.append(Document(
                page_content=section,
                metadata={"source": name, "page": i, "format": "markdown"},
            ))
    return docs or [Document(page_content=raw, metadata={"source": name, "page": 1, "format": "markdown"})]


def _load_text(file) -> list[Document]:
    """Load plain text file as a single Document."""
    raw = file.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")
    name = getattr(file, "name", "file.txt")
    return [Document(page_content=raw.strip(), metadata={"source": name, "page": 1, "format": "text"})]


_LOADERS: dict[str, callable] = {
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm":  _load_html,
    ".md":   _load_markdown,
    ".txt":  _load_text,
}


def load_document(file) -> list[Document]:
    """
    Auto-detect file type by extension and load via the matching adapter.

    Args:
        file: File-like object with a .name attribute.

    Returns:
        List of LangChain Documents.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = Path(getattr(file, "name", "")).suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        supported = ", ".join(_LOADERS.keys())
        raise ValueError(f"Unsupported format '{ext}'. Supported: {supported}")
    return loader(file)


def supported_extensions() -> list[str]:
    """Return the list of supported file extensions."""
    return list(_LOADERS.keys())
